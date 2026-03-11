"""Plain-text extraction for non-text file formats (PDF, DOCX).

Used by both the semantic indexer and the /search backend so that PDF and DOCX
files can be searched and embedded without any pre-processing step.
"""
from __future__ import annotations

from pathlib import Path

# Extensions handled by this module; all others are read as plain UTF-8.
EXTRACTABLE_EXTENSIONS = {".pdf", ".docx"}


def extract_text(file_path: Path) -> str:
    """Return the plain-text content of *file_path*.

    PDF and DOCX files are extracted via their respective libraries.
    All other files are read as UTF-8 (with replacement for bad bytes).
    Returns an empty string if extraction fails (encrypted, corrupt, etc.).
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    return file_path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(file_path: Path) -> str:
    import pypdf

    try:
        reader = pypdf.PdfReader(file_path)
        if reader.is_encrypted:
            return ""
        return "\n".join(
            text
            for page in reader.pages
            if (text := page.extract_text())
        )
    except Exception:
        return ""


def _extract_docx(file_path: Path) -> str:
    import docx

    try:
        doc = docx.Document(file_path)
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        # Tables can hold credential lists, PII etc. — include cell text too.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""
